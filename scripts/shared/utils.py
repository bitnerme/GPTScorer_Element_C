import time
import openai
import os
import backoff
from datetime import datetime
import pythoncom
import win32com.client
from pdf2image import convert_from_path
import pytesseract
import docx
import subprocess
import traceback
import json
from pathlib import Path
import shutil
import subprocess
import tempfile
from pathlib import Path

SOFFICE_PATH = r"C:\Program Files\LibreOffice\program\soffice.exe"

# IMPORTANT:
# Native DOCX extraction was found to contain an indentation defect that
# caused zero-length extraction. The defect has been corrected, but the
# v2.0 scorer intentionally bypasses native DOCX extraction because all
# production baselines were established using the DOCX→PDF→OCR workflow.
# Changing this flag requires full regression testing and baseline refresh.

USE_NATIVE_DOCX_EXTRACTION = False  # v2.0 baseline lock.  

def convert_docx_to_pdf(filepath):
    filepath = Path(filepath)

    with tempfile.TemporaryDirectory() as tmpdir:
        tmpdir = Path(tmpdir)

        if not Path(SOFFICE_PATH).exists():
            raise FileNotFoundError(f"LibreOffice not found: {SOFFICE_PATH}")

        cmd = [
            SOFFICE_PATH,
            "--headless",
            "--convert-to",
            "pdf",
            "--outdir",
            str(tmpdir),
            str(filepath),
        ]

        subprocess.run(
            cmd,
            check=True,
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE,
            text=True,
        )

        pdf_path = tmpdir / f"{filepath.stem}.pdf"

        if not pdf_path.exists():
            raise FileNotFoundError(f"DOCX to PDF conversion failed: {pdf_path}")

        final_pdf = tempfile.NamedTemporaryFile(suffix=".pdf", delete=False)
        final_pdf.close()

        shutil.copyfile(pdf_path, final_pdf.name)
        return final_pdf.name

def configure_tesseract():
    # 1. Env var override
    tesseract_path = os.environ.get("TESSERACT_PATH")
    if tesseract_path and os.path.exists(tesseract_path):
        pytesseract.pytesseract.tesseract_cmd = tesseract_path
        return

    # 2. System path (Mac/Linux or properly configured Windows)
    detected = shutil.which("tesseract")
    if detected:
        pytesseract.pytesseract.tesseract_cmd = detected
        return

    # 3. Windows fallback (YOU MUST MATCH YOUR INSTALL)
    default_win = r"C:\Program Files\Tesseract-OCR\tesseract.exe"
    if os.path.exists(default_win):
        pytesseract.pytesseract.tesseract_cmd = default_win
        return

    print("⚠️ Tesseract not found — OCR will fail")

configure_tesseract()


# scripts/shared/utils.py

def get_blended_model(element, mode):
    element = element.upper()
    mode = mode.lower()

    mapping = {
        "A": ("v1.0", "v1.2"),
        "B": ("v1.2", "v1.4b"),
        "C": ("v1.13", "v1.15"),
        "D": ("v1.8d", "v2.0"),
        "E": ("v1.2", "v1.7r"),
        "F": ("v1.2", "v1.62"),
        "G": ("v1.41", "v1.42"),
        "H": ("v1.0", "v2.0"),
        "I": ("v1.1", "v1.2"),
        "J": ("v1.0", "v1.5"),
        "K": ("v1.0", "v1.2"),
        "L": ("v1.5", "v1.6")
    }

    if element not in mapping:
        raise ValueError(f"Unknown element '{element}' in get_blended_model")

    legacy, current = mapping[element]
    return legacy if mode == "legacy" else current

def normalize_columns(df, element):
    for k in range(1,7):

        target = f"{element}{k}"

        if target not in df.columns and f"_{k}" in df.columns:
            df[target] = df[f"_{k}"]

        if target not in df.columns and f"_{k}_final" in df.columns:
            df[target] = df[f"_{k}_final"]

        # cross-element fallback
        for other in ["A", "B", "C", "D", "E", "F", "G", "H", "I", "J", "K", "L"]:
            if target not in df.columns and f"{other}{k}" in df.columns:
                df[target] = df[f"{other}{k}"]

    return df 

def check_drift(current_metrics, baseline_file):
    print("ENTER check_drift")
    print("CHECK_DRIFT FUNCTION ID:", id(check_drift))

    baseline_path = Path(baseline_file)

    if not baseline_path.exists():
        return {
            "status": "NO BASELINE",
            "message": f"Baseline file not found: {baseline_file}"
        }

    with open(baseline_path, "r") as f:
        baseline = json.load(f)

    report = {}
    failures = []

    # -----------------------------
    # Compute metric differences
    # -----------------------------

    api_mean_diff = abs(current_metrics["api_mean"] - baseline["api_mean"])
    api_std_diff = abs(current_metrics["api_std"] - baseline["api_std"])

    final_mean_diff = abs(current_metrics["final_mean"] - baseline["final_mean"])
    final_std_diff = abs(current_metrics["final_std"] - baseline["final_std"])

    report["api_mean_diff"] = api_mean_diff
    report["api_std_diff"] = api_std_diff
    report["final_mean_diff"] = final_mean_diff
    report["final_std_diff"] = final_std_diff

    # -----------------------------
    # Thresholds
    # -----------------------------

    API_MEAN_THRESHOLD = 0.25
    API_STD_THRESHOLD = 0.20
    FINAL_MEAN_THRESHOLD = 0.25
    FINAL_STD_THRESHOLD = 0.20

    if api_mean_diff > API_MEAN_THRESHOLD:
        failures.append("api_mean_shift")

    if api_std_diff > API_STD_THRESHOLD:
        failures.append("api_std_shift")

    if final_mean_diff > FINAL_MEAN_THRESHOLD:
        failures.append("final_mean_shift")

    if final_std_diff > FINAL_STD_THRESHOLD:
        failures.append("final_std_shift")

    sample_warning = None

    baseline_n = baseline.get("sample_size")
    current_n = current_metrics.get("sample_size")

    print("DEBUG sample sizes:", current_n, baseline_n)

    # -----------------------------
    # Sample size guard (NEW)
    # -----------------------------
    MIN_SAMPLE_RATIO = 0.5

    if baseline_n and current_n:
        if current_n < baseline_n * MIN_SAMPLE_RATIO:
            return {
                "status": "LOW SAMPLE",
                "sample_size_current": current_n,
                "sample_size_baseline": baseline_n,
                "sample_warning": (
                    f"Sample size too small: {current_n} vs baseline {baseline_n}. "
                    "Drift detection skipped due to low statistical power."
                ),
                "diagnostic_interpretation": (
                    "Drift check not reliable due to insufficient sample size."
                ),
                "failures": [],
                "report": report
            }

    # -----------------------------
    # Final result
    # -----------------------------

    if failures:
        status = "DRIFT DETECTED"
    else:
        status = "PASS"
   
    print("DEBUG failures:", failures)
    print("DEBUG diffs:", api_mean_diff, api_std_diff, final_mean_diff, final_std_diff)

    # -----------------------------
    # Diagnostic interpretation (NEW)
    # -----------------------------
    diagnostic_interpretation = None  # handled in controller

    print("DEBUG check_drift RETURN:", {
        "status": status,
        "failures": failures,
        "sample_size_current": current_n,
    })

    return {
        "status": status,
        "sample_size_current": current_n,
        "sample_size_baseline": baseline_n,
        "sample_warning": sample_warning,
        "failures": failures,
        "diagnostic_interpretation": None,  # handled in controller
        "report": report
    }

# --- Extract text from .docx/doc or .pdf ---
def extract_text_from_file(filepath):

    ext = os.path.splitext(filepath)[1]

    ext = ext.lower()

    if ext in (".docx", ".doc"):
        print("→ DOCX BRANCH")
        return extract_text_from_docx(filepath)
    elif ext == ".pdf":
        print("→ PDF BRANCH")
        return extract_text_from_pdf(filepath)
    else:
        raise ValueError(f"Unsupported file format: {repr(ext)}")

def get_poppler_path():
    # 1. Environment variable (optional override)
    poppler_path = os.environ.get("POPPLER_PATH")
    if poppler_path and os.path.exists(poppler_path):
        return poppler_path

    # 2. If installed system-wide (Mac/Linux via brew)
    if shutil.which("pdftoppm"):
        return None  # Let pdf2image use system path

    # 3. Windows fallback
    default_win = r"C:\poppler\poppler-25.12.0\Library\bin"
    if os.path.exists(default_win):
        return default_win

    print("⚠️ Poppler not found — PDF OCR may fail")
    return None

def run_ocr(filepath):
    pythoncom.CoInitialize()

    try:
        extracted_text = ""

        ext = os.path.splitext(filepath)[1].lower()
        abs_path = os.path.abspath(filepath)
        base, _ = os.path.splitext(abs_path)
        pdf_path = base + "_temp_ocr.pdf"

        # Kill any orphaned Word processes (optional but stabilizing)
        subprocess.run(
            ["taskkill", "/f", "/im", "WINWORD.EXE"],
            stdout=subprocess.DEVNULL,
            stderr=subprocess.DEVNULL
        )

        word = win32com.client.DispatchEx("Word.Application")
        word.Visible = False
        word.DisplayAlerts = 0

        doc = word.Documents.Open(abs_path, ReadOnly=True)
        doc.SaveAs2(pdf_path, FileFormat=17)
        doc.Close(False)

        word.Quit()
        del word

        time.sleep(0.5)  # allow COM to release

        pages = convert_from_path(
            pdf_path,
            poppler_path=get_poppler_path()
        )

        for page in pages:
            extracted_text += pytesseract.image_to_string(page) + "\n"

        if os.path.exists(pdf_path):
            os.remove(pdf_path)

        return extracted_text

    except Exception:
        print("\nFULL TRACEBACK inside run_ocr:")
        traceback.print_exc()
        raise

    finally:
        pythoncom.CoUninitialize()

def extract_text_with_fallback(filepath, min_length=50, return_metadata=False):
    text = extract_text_from_file(filepath)
    text = (text or "").strip()
    
    text = extract_text_from_file(filepath)
    text = (text or "").strip()

    metadata = {
        "ocr_used": False,
        "initial_text_length": len(text),
        "ocr_text_length": 0,
        "final_text_length": len(text),
        "extraction_method": "direct"
    }

    print(f"OCR CHECK: {filepath}")
    print(f"Original Extracted text length: {len(text)}")

    if len(text) < min_length:
        try:
            print(f"OCR triggered for {filepath}")
            ext = os.path.splitext(filepath)[1].lower()

            if ext == ".pdf":
                print("OCR: starting PDF render")
                images = convert_from_path(filepath)

            elif ext == ".docx":
                print("OCR: converting DOCX to PDF")
                pdf_path = convert_docx_to_pdf(filepath)
                try:
                    print("OCR: starting converted PDF render")
                    images = convert_from_path(pdf_path)
                finally:
                    if os.path.exists(pdf_path):
                        os.remove(pdf_path)
            else:
                print(f"OCR skipped for unsupported file type: {ext}")
                return (text, metadata) if return_metadata else text

            print(f"OCR: completed PDF render, pages={len(images)}")

            ocr_pages = []
            for idx, image in enumerate(images, start=1):
                print(f"OCR: starting page {idx}")
                page_text = pytesseract.image_to_string(image)
                print(f"OCR: completed page {idx}, chars={len(page_text)}")
                ocr_pages.append(page_text)

            ocr_text = "\n".join(ocr_pages).strip()

            metadata.update({
                "ocr_used": True,
                "ocr_text_length": len(ocr_text),
                "final_text_length": len(ocr_text),
                "extraction_method": "ocr_fallback"
            })

            print(f"OCR returned length: {len(ocr_text)}")

            return (ocr_text, metadata) if return_metadata else ocr_text
            

        except Exception as e:
            print(f"❌ OCR FAILED for {filepath}: {e}")
            traceback.print_exc()
            metadata["extraction_method"] = "ocr_failed"
            metadata["final_text_length"] = 0
            return ("", metadata) if return_metadata else ""

    return (text, metadata) if return_metadata else text

def extract_text_from_docx(filepath):

    # v2.0 policy:
    # Native DOCX extraction has been repaired but is intentionally bypassed.
    # Regression testing showed v2.0 baselines were established using the
    # DOCX → PDF render/OCR extraction path. Enabling native DOCX extraction
    # would change the scored text representation and require new baselines.

    if not USE_NATIVE_DOCX_EXTRACTION:
        print("DOCX native extraction intentionally bypassed; returning 0 chars")
        return ""
        
    abs_path = os.path.abspath(filepath)
    doc = docx.Document(abs_path)
    parts = []

    def add_paragraphs(paragraphs):
        for para in paragraphs:
            text = para.text.strip()
            if text:
                parts.append(text)

    def add_table(table):
        for row in table.rows:
            for cell in row.cells:
                add_paragraphs(cell.paragraphs)
                for nested_table in cell.tables:
                    add_table(nested_table)

    add_paragraphs(doc.paragraphs)

    for table in doc.tables:
        add_table(table)

    for section in doc.sections:
        add_paragraphs(section.header.paragraphs)
        add_paragraphs(section.footer.paragraphs)
        for table in section.header.tables:
            add_table(table)
        for table in section.footer.tables:
            add_table(table)

    return "\n".join(parts)
    
def extract_text_from_pdf(filepath):
    try:
        from pdfminer.high_level import extract_text
    except ImportError:
        raise ImportError("pdfminer.six must be installed to extract text from PDFs")

    text = extract_text(filepath)
    return text

# --- GPT call with retries ---
@backoff.on_exception(backoff.expo, openai.error.OpenAIError, max_tries=5)
def call_gpt_with_backoff(prompt, system="You are a helpful assistant.",
                          model_order=None, temperature=0.0, max_tokens=3500):

    if model_order is None:
        raise ValueError("model_order must be explicitly provided by caller")

    last_exception = None

    for current_model in model_order:
        try:
            print(f"🔄 Trying model: {current_model}")

            response = openai.ChatCompletion.create(
                model=current_model,
                messages=[
                    {"role": "system", "content": system},
                    {"role": "user", "content": prompt}
                ],
                temperature=temperature,
                max_tokens=max_tokens,
            )

            content = response['choices'][0]['message']['content']
            
            print("RAW GPT RESPONSE:", response)

            return content

        except Exception as e:
            print(f"⚠️ Error with model {current_model}: {e}")
            last_exception = e
            time.sleep(1)

    print("❌ All GPT models failed.")
    raise last_exception


